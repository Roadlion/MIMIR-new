import io
import re

def markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """
    Very basic markdown to docx converter.
    Handles headers, bold, and basic paragraphs.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    except ImportError:
        raise RuntimeError("python-docx is not installed. Please install python-docx to export DOCX files.")

    doc = Document()
    
    # Title
    title = doc.add_heading("MIMIR Oracle Research Report", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            # Handle bold within list items
            _process_inline_formatting(p, line[2:])
        else:
            p = doc.add_paragraph()
            _process_inline_formatting(p, line)
            
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def _process_inline_formatting(paragraph, text):
    """Helper to process **bold** and *italic* text in a paragraph."""
    # This is a very simplified parser
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
